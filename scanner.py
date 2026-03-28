#!/usr/bin/env python3
"""
PYT Radar â Automated podcast & YouTube episode scanner.
Runs as a GitHub Action on a cron schedule.

Searches for new episodes, transcribes via AssemblyAI,
generates summary + transcript docs, uploads to Dropbox,
and updates the GitHub Pages blog + RSS feed.
"""

import os
import sys
import json
import time
import re
import shutil
import subprocess
import tempfile
import argparse
import requests
from datetime import datetime, timedelta, timezone
from pathlib import Path

try:
    from youtube_transcript_api import YouTubeTranscriptApi
    from youtube_transcript_api.proxies import WebshareProxyConfig
    HAS_YT_TRANSCRIPT = True
except ImportError:
    HAS_YT_TRANSCRIPT = False


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# Configuration
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

ASSEMBLYAI_KEY = os.environ["ASSEMBLYAI_API_KEY"]
DROPBOX_REFRESH_TOKEN = os.environ.get("DROPBOX_REFRESH_TOKEN", "")
DROPBOX_APP_KEY = os.environ.get("DROPBOX_APP_KEY", "")
DROPBOX_APP_SECRET = os.environ.get("DROPBOX_APP_SECRET", "")
DROPBOX_TOKEN = os.environ.get("DROPBOX_TOKEN", "")  # fallback; prefer refresh token
YOUTUBE_API_KEY = os.environ.get("YOUTUBE_API_KEY", "")
LISTENNOTES_API_KEY = os.environ.get("LISTENNOTES_API_KEY", "")

SEARCH_TARGETS = json.loads(os.environ.get("SEARCH_TARGETS", '{"guests":["Karen Bass"],"topics":["Nithya Raman"]}'))
BLOG_URL = "https://ericspiegelman.github.io/pyt-radar/"
REPO_ROOT = Path(os.environ.get("GITHUB_WORKSPACE", "."))
EPISODES_FILE = REPO_ROOT / "data" / "found_episodes.json"
INDEX_FILE = REPO_ROOT / "index.html"
FEED_FILE = REPO_ROOT / "feed.xml"
KB_DIR = REPO_ROOT / "kb"
COOKIES_FILE = Path("cookies.txt")
WEBSHARE_PROXY_USERNAME = os.environ.get("WEBSHARE_PROXY_USERNAME", "")
WEBSHARE_PROXY_PASSWORD = os.environ.get("WEBSHARE_PROXY_PASSWORD", "")

# Global flag for YouTube cookie expiry detection
_youtube_cookies_expired = False


# ────────────────────────────────────────────────────────
# Dropbox token refresh
# ────────────────────────────────────────────────────────

def refresh_dropbox_token():
    """Use the long-lived refresh token to get a fresh short-lived access token.
    Falls back to DROPBOX_TOKEN env var if refresh token is not configured."""
    global DROPBOX_TOKEN
    if not DROPBOX_REFRESH_TOKEN or not DROPBOX_APP_KEY or not DROPBOX_APP_SECRET:
        if DROPBOX_TOKEN:
            print("  Dropbox: using static access token (may be expired)")
        else:
            print("  Dropbox: no credentials configured - uploads will be skipped")
        return
    try:
        resp = requests.post(
            "https://api.dropboxapi.com/oauth2/token",
            data={
                "grant_type": "refresh_token",
                "refresh_token": DROPBOX_REFRESH_TOKEN,
                "client_id": DROPBOX_APP_KEY,
                "client_secret": DROPBOX_APP_SECRET,
            },
        )
        resp.raise_for_status()
        DROPBOX_TOKEN = resp.json()["access_token"]
        print("  Dropbox: refreshed access token successfully")
    except Exception as e:
        print(f"  WARNING: Dropbox token refresh failed: {e}")
        if DROPBOX_TOKEN:
            print("  Falling back to static DROPBOX_TOKEN")

# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# Episode log (deduplication)
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

def load_episodes():
    if EPISODES_FILE.exists():
        return json.loads(EPISODES_FILE.read_text())
    return {"episodes": []}


def save_episodes(data):
    EPISODES_FILE.parent.mkdir(parents=True, exist_ok=True)
    EPISODES_FILE.write_text(json.dumps(data, indent=2))


def is_duplicate(url, episodes_data):
    return any(ep["url"] == url for ep in episodes_data["episodes"])


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# YouTube Data API search
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

def search_youtube(query, max_results=10, published_after=None):
    """Search YouTube using the Data API v3."""
    url = "https://www.googleapis.com/youtube/v3/search"
    params = {
        "part": "snippet",
        "q": query,
        "type": "video",
        "maxResults": max_results,
        "order": "date",
        "key": YOUTUBE_API_KEY,
    }
    if published_after:
        params["publishedAfter"] = published_after.strftime("%Y-%m-%dT%H:%M:%SZ")

    resp = requests.get(url, params=params)
    if resp.status_code == 403:
        print(f"  WARNING: YouTube API quota exceeded or forbidden (403). Skipping this query.")
        return []
    resp.raise_for_status()
    data = resp.json()

    results = []
    for item in data.get("items", []):
        vid = item["id"].get("videoId")
        if not vid:
            continue
        snippet = item["snippet"]
        results.append({
            "video_id": vid,
            "url": f"https://www.youtube.com/watch?v={vid}",
            "title": snippet["title"],
            "channel": snippet["channelTitle"],
            "published": snippet["publishedAt"],
            "description": snippet.get("description", ""),
        })
    return results


def get_video_duration(video_id):
    """Get video duration in seconds via YouTube Data API."""
    url = "https://www.googleapis.com/youtube/v3/videos"
    params = {
        "part": "contentDetails",
        "id": video_id,
        "key": YOUTUBE_API_KEY,
    }
    resp = requests.get(url, params=params)
    if resp.status_code == 403:
        print(f"  WARNING: YouTube API quota exceeded (403) for video duration check")
        return 0
    resp.raise_for_status()
    items = resp.json().get("items", [])
    if not items:
        return 0
    duration_str = items[0]["contentDetails"]["duration"]  # e.g. PT1H23M45S
    hours = int(re.search(r'(\d+)H', duration_str).group(1)) if 'H' in duration_str else 0
    minutes = int(re.search(r'(\d+)M', duration_str).group(1)) if 'M' in duration_str else 0
    seconds = int(re.search(r'(\d+)S', duration_str).group(1)) if 'S' in duration_str else 0
    return hours * 3600 + minutes * 60 + seconds


def parse_target(target):
    """Parse a search target which can be a string or {"name": ..., "context": ...}."""
    if isinstance(target, dict):
        return target["name"], target.get("context", "")
    return target, ""


def find_youtube_episodes(episodes_data):
    """Search YouTube for new episodes matching all targets."""
    new_episodes = []
    cutoff = datetime.now(timezone.utc) - timedelta(days=7)

    for target_entry in SEARCH_TARGETS.get("guests", []):
        guest, context = parse_target(target_entry)
        print(f"Searching YouTube for guest: {guest}")
        ctx = f" {context}" if context else ""
        queries = [
            f'"{guest}"{ctx} interview OR podcast OR guest',
            f'"{guest}"{ctx} podcast episode',
        ]
        for q in queries:
            results = search_youtube(q, max_results=10, published_after=cutoff)
            for r in results:
                if is_duplicate(r["url"], episodes_data):
                    print(f"  Skip (duplicate): {r['title']}")
                    continue
                if any(e["url"] == r["url"] for e in new_episodes):
                    continue
                duration = get_video_duration(r["video_id"])
                if duration < 300:
                    print(f"  Skip (too short {duration}s): {r['title']}")
                    continue
                print(f"  FOUND: {r['title']} ({r['channel']})")
                new_episodes.append({
                    "url": r["url"],
                    "video_id": r["video_id"],
                    "title": r["title"],
                    "show_name": r["channel"],
                    "date_published": r["published"][:10],
                    "search_target": guest,
                    "match_type": "guest",
                    "duration": duration,
                })

    for target_entry in SEARCH_TARGETS.get("topics", []):
        topic, context = parse_target(target_entry)
        print(f"Searching YouTube for topic: {topic}")
        ctx = f" {context}" if context else ""
        queries = [
            f'"{topic}"{ctx} interview OR podcast OR discussion',
            f'"{topic}"{ctx} podcast episode',
        ]
        for q in queries:
            results = search_youtube(q, max_results=10, published_after=cutoff)
            for r in results:
                if is_duplicate(r["url"], episodes_data):
                    print(f"  Skip (duplicate): {r['title']}")
                    continue
                if any(e["url"] == r["url"] for e in new_episodes):
                    continue
                duration = get_video_duration(r["video_id"])
                if duration < 300:
                    print(f"  Skip (too short {duration}s): {r['title']}")
                    continue
                print(f"  FOUND: {r['title']} ({r['channel']})")
                new_episodes.append({
                    "url": r["url"],
                    "video_id": r["video_id"],
                    "title": r["title"],
                    "show_name": r["channel"],
                    "date_published": r["published"][:10],
                    "search_target": topic,
                    "match_type": "mentioned",
                    "duration": duration,
                })

    return new_episodes


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# Listen Notes API search (podcasts)
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

def search_listen_notes(query, published_after=None):
    """Search Listen Notes for podcast episodes matching a query."""
    if not LISTENNOTES_API_KEY:
        print("  WARNING: No LISTENNOTES_API_KEY â skipping podcast search")
        return []

    url = "https://listen-api.listennotes.com/api/v2/search"
    headers = {"X-ListenAPI-Key": LISTENNOTES_API_KEY}
    params = {
        "q": query,
        "type": "episode",
        "sort_by_date": 1,  # sort by date (newest first)
        "len_min": 5,  # at least 5 minutes
    }
    if published_after:
        params["published_after"] = int(published_after.timestamp() * 1000)

    try:
        resp = requests.get(url, headers=headers, params=params)
        resp.raise_for_status()
    except requests.exceptions.HTTPError as e:
        print(f"  Listen Notes API error: {e}")
        return []

    data = resp.json()
    results = []
    for item in data.get("results", []):
        # Build a URL â prefer the Listen Notes link, but also store the audio
        ln_url = item.get("listennotes_url", "")
        audio_url = item.get("audio", "")
        episode_url = item.get("link", ln_url)  # original episode link if available

        results.append({
            "url": episode_url or ln_url,
            "listennotes_url": ln_url,
            "audio_url": audio_url,
            "title": item.get("title_original", ""),
            "show_name": item.get("podcast", {}).get("title_original", "Unknown"),
            "published": item.get("pub_date_ms", 0),
            "description": item.get("description_original", ""),
            "duration": (item.get("audio_length_sec") or 0),
        })
    return results


def find_podcast_episodes(episodes_data):
    """Search Listen Notes for new podcast episodes matching all targets."""
    new_episodes = []
    cutoff = datetime.now(timezone.utc) - timedelta(days=7)

    for target_entry in SEARCH_TARGETS.get("guests", []):
        guest, context = parse_target(target_entry)
        print(f"Searching podcasts for guest: {guest}")
        query = f'"{guest}"' + (f" {context}" if context else "")
        results = search_listen_notes(query, published_after=cutoff)
        for r in results:
            if is_duplicate(r["url"], episodes_data):
                print(f"  Skip (duplicate): {r['title']}")
                continue
            if any(e["url"] == r["url"] for e in new_episodes):
                continue
            if r["duration"] < 300:
                print(f"  Skip (too short {r['duration']}s): {r['title']}")
                continue
            print(f"  FOUND: {r['title']} ({r['show_name']})")
            new_episodes.append({
                "url": r["url"],
                "video_id": None,
                "audio_url": r["audio_url"],
                "title": r["title"],
                "show_name": r["show_name"],
                "date_published": datetime.fromtimestamp(
                    r["published"] / 1000, tz=timezone.utc
                ).strftime("%Y-%m-%d") if r["published"] else "unknown",
                "search_target": guest,
                "match_type": "guest",
                "duration": r["duration"],
            })

    for target_entry in SEARCH_TARGETS.get("topics", []):
        topic, context = parse_target(target_entry)
        print(f"Searching podcasts for topic: {topic}")
        query = f'"{topic}"' + (f" {context}" if context else "")
        results = search_listen_notes(query, published_after=cutoff)
        for r in results:
            if is_duplicate(r["url"], episodes_data):
                print(f"  Skip (duplicate): {r['title']}")
                continue
            if any(e["url"] == r["url"] for e in new_episodes):
                continue
            if r["duration"] < 300:
                print(f"  Skip (too short {r['duration']}s): {r['title']}")
                continue
            print(f"  FOUND: {r['title']} ({r['show_name']})")
            new_episodes.append({
                "url": r["url"],
                "video_id": None,
                "audio_url": r["audio_url"],
                "title": r["title"],
                "show_name": r["show_name"],
                "date_published": datetime.fromtimestamp(
                    r["published"] / 1000, tz=timezone.utc
                ).strftime("%Y-%m-%d") if r["published"] else "unknown",
                "search_target": topic,
                "match_type": "mentioned",
                "duration": r["duration"],
            })

    return new_episodes


def find_new_episodes(episodes_data, mode="all"):
    """Search for new episodes. mode: 'youtube', 'podcast', or 'all'."""
    new_episodes = []
    if mode in ("youtube", "all"):
        new_episodes.extend(find_youtube_episodes(episodes_data))
    if mode in ("podcast", "all"):
        new_episodes.extend(find_podcast_episodes(episodes_data))
    return new_episodes


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# Transcription (AssemblyAI)
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

def download_youtube_audio(url):
    """Download audio from YouTube using yt-dlp with resilient options.
    Returns local file path. Tries multiple strategies to work around
    YouTube blocking datacenter IPs (common on GitHub Actions runners)."""
    tmpdir = tempfile.mkdtemp()
    output_template = os.path.join(tmpdir, "audio.%(ext)s")

    # Strategy 1: Use iOS player client (less aggressive bot detection)
    cookie_args = ["--cookies", str(COOKIES_FILE)] if COOKIES_FILE.exists() else []
    if cookie_args:
        print("  Using YouTube cookies for authentication")

    strategies = [
        {
            "name": "iOS client",
            "args": [
                "yt-dlp",
                "--js-runtimes", "node",
                "-f", "bestaudio/best",
                "--extractor-args", "youtube:player_client=web_creator",
                "--no-check-certificates",
                "--retries", "3",
                "--socket-timeout", "30",
                "--user-agent", "Mozilla/5.0 (iPhone; CPU iPhone OS 16_0 like Mac OS X) AppleWebKit/605.1.15",
                "-o", output_template,
                url,
            ],
        },
        {
            "name": "web + mweb clients",
            "args": [
                "yt-dlp",
                "--js-runtimes", "node",
                "-f", "bestaudio/best",
                "--extractor-args", "youtube:player_client=tv_embedded",
                "--no-check-certificates",
                "--retries", "3",
                "--socket-timeout", "30",
                "-o", output_template,
                url,
            ],
        },
        {
            "name": "default with fallback format",
            "args": [
                "yt-dlp",
                "--js-runtimes", "node",
                "-f", "worstaudio/worst",
                "--no-check-certificates",
                "--retries", "3",
                "--socket-timeout", "30",
                "-o", output_template,
                url,
            ],
        },
    ]

    last_error = None
    for strategy in strategies:
        # Clean any leftover files from previous attempts
        for f in os.listdir(tmpdir):
            os.unlink(os.path.join(tmpdir, f))
        try:
            print(f"  yt-dlp strategy: {strategy['name']}")
            full_args = strategy["args"][:-1] + cookie_args + [strategy["args"][-1]]
            result = subprocess.run(
                full_args,
                capture_output=True, text=True, timeout=120
            )
            if result.returncode == 0:
                for f in os.listdir(tmpdir):
                    if f.startswith("audio."):
                        print(f"  yt-dlp succeeded with strategy: {strategy['name']}")
                        return os.path.join(tmpdir, f)
            else:
                print(f"  yt-dlp failed ({strategy['name']}): {result.stderr[:200]}")
                last_error = result.stderr
                if "cookies are no longer valid" in result.stderr or "Sign in to confirm" in result.stderr:
                    global _youtube_cookies_expired
                    _youtube_cookies_expired = True
        except subprocess.TimeoutExpired:
            print(f"  yt-dlp timed out ({strategy['name']})")
            last_error = "timeout"
        except Exception as e:
            print(f"  yt-dlp error ({strategy['name']}): {e}")
            last_error = str(e)

    # Clean up tmpdir on failure
    shutil.rmtree(tmpdir, ignore_errors=True)
    raise RuntimeError(f"All yt-dlp strategies failed for {url}. Last error: {last_error}")


def get_youtube_audio_url(url):
    """Try to extract the direct audio stream URL from YouTube using yt-dlp -g.
    Returns the URL string if successful, None otherwise.
    This URL can be passed directly to AssemblyAI for transcription."""
    cookie_args = ["--cookies", str(COOKIES_FILE)] if COOKIES_FILE.exists() else []
    strategies = [
        ["yt-dlp", "--js-runtimes", "node", "-g", "-f", "bestaudio/best",
         "--extractor-args", "youtube:player_client=web_creator",
         "--no-check-certificates", url],
        ["yt-dlp", "--js-runtimes", "node", "-g", "-f", "bestaudio/best",
         "--extractor-args", "youtube:player_client=tv_embedded",
         "--no-check-certificates", url],
        ["yt-dlp", "--js-runtimes", "node", "-g", "-f", "worstaudio/worst",
         "--no-check-certificates", url],
    ]
    for args in strategies:
        try:
            full_args = args[:-1] + cookie_args + [args[-1]]
            result = subprocess.run(full_args, capture_output=True, text=True, timeout=60)
            if result.returncode == 0 and result.stdout.strip():
                audio_url = result.stdout.strip().split("\n")[0]
                print(f"  Got direct audio URL via yt-dlp -g")
                return audio_url
            elif result.stderr and ("cookies are no longer valid" in result.stderr or "Sign in to confirm" in result.stderr):
                global _youtube_cookies_expired
                _youtube_cookies_expired = True
        except Exception:
            continue
    return None


def upload_to_assemblyai(filepath):
    """Upload a local audio file to AssemblyAI and return the upload URL."""
    print("  Uploading audio to AssemblyAI...")
    with open(filepath, "rb") as f:
        resp = requests.post(
            "https://api.assemblyai.com/v2/upload",
            headers={"Authorization": ASSEMBLYAI_KEY},
            data=f
        )
    resp.raise_for_status()
    return resp.json()["upload_url"]


def transcribe(audio_url):
    """Submit transcription and poll until complete. Returns transcript dict."""
    print("  Submitting transcription request...")
    resp = requests.post(
        "https://api.assemblyai.com/v2/transcript",
        headers={
            "Authorization": ASSEMBLYAI_KEY,
            "Content-Type": "application/json",
        },
        json={
            "audio_url": audio_url,
            "speaker_labels": True,
            "auto_highlights": True,
            "speech_models": ["universal-3-pro"],
        }
    )
    resp.raise_for_status()
    transcript_id = resp.json()["id"]
    print(f"  Transcript ID: {transcript_id}")

    while True:
        resp = requests.get(
            f"https://api.assemblyai.com/v2/transcript/{transcript_id}",
            headers={"Authorization": ASSEMBLYAI_KEY}
        )
        resp.raise_for_status()
        data = resp.json()
        status = data["status"]

        if status == "completed":
            print("  Transcription complete!")
            return data
        elif status == "error":
            raise RuntimeError(f"Transcription failed: {data.get('error', 'Unknown')}")

        print(f"  Status: {status} â waiting 15s...")
        time.sleep(15)


def get_youtube_transcript(video_id):
    """Try to get YouTube captions/transcript via youtube-transcript-api.
    Returns transcript_data dict compatible with AssemblyAI format, or None."""
    if not HAS_YT_TRANSCRIPT:
        print("  youtube-transcript-api not available")
        return None
    try:
        proxy_config = None
        if WEBSHARE_PROXY_USERNAME and WEBSHARE_PROXY_PASSWORD:
            proxy_config = WebshareProxyConfig(
                proxy_username=WEBSHARE_PROXY_USERNAME,
                proxy_password=WEBSHARE_PROXY_PASSWORD,
            )
            print("  Using Webshare proxy for YouTube captions")
        ytt_api = YouTubeTranscriptApi(proxy_config=proxy_config)
        transcript_list = ytt_api.fetch(video_id)
        # Convert to AssemblyAI-compatible format
        utterances = []
        full_text_parts = []
        for entry in transcript_list:
            start_ms = int(entry.start * 1000)
            end_ms = int((entry.start + entry.duration) * 1000)
            text = entry.text.strip()
            if not text:
                continue
            utterances.append({
                "start": start_ms,
                "end": end_ms,
                "text": text,
                "speaker": "A",  # captions don't have speaker labels
            })
            full_text_parts.append(text)
        if not utterances:
            print("  YouTube captions were empty")
            return None
        total_duration = utterances[-1]["end"] // 1000 if utterances else 0
        print(f"  Got YouTube captions: {len(utterances)} segments, ~{total_duration // 60}m")
        return {
            "utterances": utterances,
            "text": " ".join(full_text_parts),
            "audio_duration": total_duration,
            "source": "youtube_captions",
        }
    except Exception as e:
        print(f"  YouTube captions not available: {e}")
        return None

def transcribe_episode(episode):
    """Full pipeline: get transcript for episode.
    For YouTube, tries in order:
      1. YouTube captions via youtube-transcript-api (no download needed)
      2. Download audio with yt-dlp, upload to AssemblyAI
      3. Extract direct stream URL with yt-dlp -g, pass to AssemblyAI
    For podcasts, uses direct audio URL with AssemblyAI.
    """
    print(f"Transcribing: {episode['title']}")

    if episode.get("video_id"):
        # Strategy 1: Try YouTube captions first (avoids bot detection entirely)
        transcript_data = get_youtube_transcript(episode["video_id"])
        if transcript_data:
            print("  Using YouTube captions (no audio download needed)")
            return transcript_data

        # Strategy 2: Try yt-dlp download + AssemblyAI
        audio_url_for_aai = None
        try:
            audio_path = download_youtube_audio(episode["url"])
            try:
                audio_url_for_aai = upload_to_assemblyai(audio_path)
            finally:
                os.unlink(audio_path)
                os.rmdir(os.path.dirname(audio_path))
        except RuntimeError as e:
            print(f"  Download failed, trying direct URL extraction: {e}")
            # Strategy 3: Try yt-dlp -g for direct URL
            direct_url = get_youtube_audio_url(episode["url"])
            if direct_url:
                print(f"  Using direct stream URL for AssemblyAI")
                audio_url_for_aai = direct_url

        if audio_url_for_aai:
            return transcribe(audio_url_for_aai)

        # Strategy 4: Pass YouTube URL directly to AssemblyAI
        # AssemblyAI can download from YouTube using their own infrastructure
        print(f"  Trying AssemblyAI direct YouTube URL transcription...")
        try:
            return transcribe(episode["url"])
        except Exception as e:
            print(f"  AssemblyAI direct URL also failed: {e}")

        print(f"  WARNING: All YouTube transcript strategies failed for: {episode['title']}")
        return None

    elif episode.get("audio_url"):
        print(f"  Using direct audio URL from podcast")
        return transcribe(episode["audio_url"])
    else:
        raise RuntimeError(f"No audio source for episode: {episode['title']}")

def generate_summary_with_claude(episode, transcript_data):
    """Use Claude API to generate a structured summary of the transcript."""
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("  WARNING: No ANTHROPIC_API_KEY â generating basic summary")
        return generate_basic_summary(episode, transcript_data)

    # Build a condensed version of the transcript for Claude
    utterances = transcript_data.get("utterances", [])
    transcript_text = ""
    for u in utterances:
        start_sec = u["start"] // 1000
        mm, ss = divmod(start_sec, 60)
        transcript_text += f"[{mm:02d}:{ss:02d}] Speaker {u['speaker']}: {u['text']}\n"

    # Truncate if too long (keep first and last portions)
    if len(transcript_text) > 80000:
        half = 38000
        transcript_text = transcript_text[:half] + "\n\n[...TRUNCATED...]\n\n" + transcript_text[-half:]

    prompt = f"""Analyze this podcast/YouTube transcript and generate a structured summary.

EPISODE INFO:
- Title: {episode['title']}
- Show: {episode['show_name']}
- Search target: {episode['search_target']} ({episode['match_type']})
- URL: {episode['url']}

TRANSCRIPT:
{transcript_text}

Generate a JSON response with this exact structure:
{{
  "overview": "2-3 sentence overview of what the episode is about and why the search target is relevant",
  "appearance_type": "guest|mentioned",
  "appearance_type_explanation": "Is {episode['search_target']} actually ON the show as a guest/interviewee, or are they just being discussed/mentioned by others? 'guest' means they are speaking on the episode. 'mentioned' means others are talking about them.",
  "topics": ["topic 1", "topic 2", ...],
  "sentiment": "positive|negative|neutral|mixed",
  "sentiment_explanation": "2-3 sentences explaining the tone toward the search target with specific examples",
  "key_quotes": [
    {{"text": "the quote", "speaker": "Speaker Name or Speaker X", "timestamp_mm_ss": "MM:SS", "timestamp_seconds": 1234}},
    ...
  ],
  "key_moment": {{
    "text": "the single most important quote about the search target",
    "speaker": "Speaker Name",
    "timestamp_mm_ss": "MM:SS",
    "timestamp_seconds": 1234,
    "context": "1 sentence explaining why this moment matters"
  }},

QUOTE RULES (CRITICAL):
- Every quote in key_quotes and the key_moment MUST be specifically about {episode['search_target']}. Quotes must mention {episode['search_target']} by name, refer to them directly, or be {episode['search_target']} speaking about themselves/their work.
- Do NOT include generic quotes about other topics that don't relate to {episode['search_target']}.
- Each quote MUST be 1-2 sentences max, never more than 50 words.
- Do NOT dump long transcript passages. Extract only the most pointed, illustrative sentence or two.
- Include up to 8 short quotes from different moments where {episode['search_target']} is discussed or speaking.
- The key_moment text must also be a single short quote about {episode['search_target']}, not a passage.
- If someone spoke at length about the target, pick multiple SHORT quotes from different moments rather than one long block.

  "sections": [
    {{"start": "MM:SS", "end": "MM:SS", "description": "what was discussed"}},
    ...
  ]
}}

IMPORTANT: For timestamps, provide both MM:SS format and total seconds. Identify speakers by name if possible based on context (show hosts, known guests), otherwise use Speaker A/B/C labels."""

    resp = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
        json={
            "model": "claude-sonnet-4-20250514",
            "max_tokens": 4096,
            "messages": [{"role": "user", "content": prompt}],
        }
    )
    resp.raise_for_status()
    content = resp.json()["content"][0]["text"]

    # Extract JSON from response
    json_match = re.search(r'\{[\s\S]*\}', content)
    if json_match:
        return json.loads(json_match.group())
    else:
        print("  WARNING: Could not parse Claude response, using basic summary")
        return generate_basic_summary(episode, transcript_data)


def generate_basic_summary(episode, transcript_data):
    """Fallback summary when Claude API is not available."""
    utterances = transcript_data.get("utterances", [])
    duration_sec = transcript_data.get("audio_duration", 0)
    duration_min = duration_sec // 60

    # Find utterances mentioning the search target
    target = episode["search_target"].lower()
    matching = []
    for u in utterances:
        if target in u["text"].lower():
            start_sec = u["start"] // 1000
            mm, ss = divmod(start_sec, 60)
            matching.append({
                "text": u["text"],
                "speaker": f"Speaker {u['speaker']}",
                "timestamp_mm_ss": f"{mm:02d}:{ss:02d}",
                "timestamp_seconds": start_sec,
            })

    key_moment = matching[0] if matching else {
        "text": "Search target mentioned in this episode",
        "speaker": "Unknown",
        "timestamp_mm_ss": "00:00",
        "timestamp_seconds": 0,
        "context": "Automated detection â review recommended",
    }
    if "context" not in key_moment:
        key_moment["context"] = f"{episode['search_target']} was discussed in this episode"

    return {
        "overview": f"This {duration_min}-minute episode of {episode['show_name']} mentions {episode['search_target']}.",
        "appearance_type": "mentioned",  # Default to mentioned without AI analysis
        "topics": [episode["search_target"], episode["show_name"]],
        "sentiment": "neutral",
        "sentiment_explanation": "Automated analysis â sentiment could not be determined without Claude API.",
        "key_quotes": matching[:5],
        "key_moment": key_moment,
        "sections": [{"start": "00:00", "end": f"{duration_min}:00", "description": "Full episode"}],
    }


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# Document generation (.docx)
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

def create_summary_docx(episode, summary, output_path):
    """Create a summary .docx file. Uses python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    doc.add_heading("PYT Radar â Episode Summary", level=1)
    doc.add_paragraph("")

    # Metadata
    meta_fields = [
        ("Episode", episode["title"]),
        ("Show", episode["show_name"]),
        ("Date Found", datetime.now().strftime("%B %d, %Y")),
        ("Published", episode["date_published"]),
        ("URL", episode["url"]),
        ("Search Target", f"{episode['search_target']} ({episode['match_type']})"),
    ]
    for label, value in meta_fields:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)

    doc.add_paragraph("")

    # Overview
    doc.add_heading("Overview", level=2)
    doc.add_paragraph(summary["overview"])

    # Topics
    doc.add_heading("Topics Covered", level=2)
    for topic in summary.get("topics", []):
        doc.add_paragraph(topic, style="List Bullet")

    # Sentiment
    doc.add_heading("Sentiment Analysis", level=2)
    p = doc.add_paragraph()
    run = p.add_run(f"Sentiment toward {episode['search_target']}: {summary['sentiment'].upper()}")
    run.bold = True
    doc.add_paragraph(summary.get("sentiment_explanation", ""))

    # Key Quotes
    doc.add_heading("Key Quotes", level=2)
    for i, q in enumerate(summary.get("key_quotes", []), 1):
        ts = q.get("timestamp_mm_ss", "??:??")
        secs = q.get("timestamp_seconds", 0)
        speaker = q.get("speaker", "Unknown")
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. [{ts} / {secs}s] {speaker}: ")
        run.bold = True
        p.add_run(f'"{q["text"]}"')

    # Key Moment
    km = summary.get("key_moment", {})
    if km:
        doc.add_heading("Key Moment", level=2)
        p = doc.add_paragraph()
        run = p.add_run("Timestamp: ")
        run.bold = True
        p.add_run(f"{km.get('timestamp_mm_ss', '??:??')} ({km.get('timestamp_seconds', 0)} seconds)")

        p = doc.add_paragraph()
        run = p.add_run("Speaker: ")
        run.bold = True
        p.add_run(km.get("speaker", "Unknown"))

        p = doc.add_paragraph()
        run = p.add_run("Quote: ")
        run.bold = True
        p.add_run(f'"{km.get("text", "")}"')

        yt_link = f"{episode['url']}&t={km.get('timestamp_seconds', 0)}"
        p = doc.add_paragraph()
        run = p.add_run("YouTube timecoded link: ")
        run.bold = True
        p.add_run(yt_link)

    doc.save(output_path)
    print(f"  Summary saved: {output_path}")


def create_transcript_docx(episode, transcript_data, output_path):
    """Create a transcript .docx file."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(episode["title"], level=1)
    doc.add_paragraph("")

    meta_fields = [
        ("Show", episode["show_name"]),
        ("Date Published", episode["date_published"]),
        ("Duration", f"{transcript_data.get('audio_duration', 0) // 60} minutes"),
        ("URL", episode["url"]),
        ("Transcribed by", "AssemblyAI (speaker diarization enabled)"),
    ]
    for label, value in meta_fields:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)

    doc.add_paragraph("")
    doc.add_heading("Transcript", level=2)
    doc.add_paragraph("")

    for u in transcript_data.get("utterances", []):
        start_sec = u["start"] // 1000
        mm, ss = divmod(start_sec, 60)
        p = doc.add_paragraph()
        ts_run = p.add_run(f"[{mm:02d}:{ss:02d}] ")
        ts_run.font.color.rgb = RGBColor(128, 128, 128)
        ts_run.font.size = Pt(9)
        speaker_run = p.add_run(f"Speaker {u['speaker']}: ")
        speaker_run.bold = True
        p.add_run(u["text"])

    doc.save(output_path)
    print(f"  Transcript saved: {output_path}")


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# Dropbox upload + shared links
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

def upload_to_dropbox(local_path, dropbox_path):
    """Upload a file to Dropbox."""
    print(f"  Uploading to Dropbox: {dropbox_path}")
    with open(local_path, "rb") as f:
        resp = requests.post(
            "https://content.dropboxapi.com/2/files/upload",
            headers={
                "Authorization": f"Bearer {DROPBOX_TOKEN}",
                "Dropbox-API-Arg": json.dumps({
                    "path": dropbox_path,
                    "mode": "overwrite",
                    "autorename": True,
                    "mute": False,
                }),
                "Content-Type": "application/octet-stream",
            },
            data=f,
        )
    resp.raise_for_status()
    result = resp.json()
    if "error_summary" in result:
        raise RuntimeError(f"Dropbox upload failed: {result['error_summary']}")
    print(f"  Uploaded: {result.get('path_display')}")


def get_dropbox_link(dropbox_path):
    """Get or create a shared link for a Dropbox file."""
    resp = requests.post(
        "https://api.dropboxapi.com/2/sharing/create_shared_link_with_settings",
        headers={
            "Authorization": f"Bearer {DROPBOX_TOKEN}",
            "Content-Type": "application/json",
        },
        json={
            "path": dropbox_path,
            "settings": {"requested_visibility": "public"},
        }
    )
    data = resp.json()
    error = data.get("error_summary", "")

    if "shared_link_already_exists" in error:
        resp2 = requests.post(
            "https://api.dropboxapi.com/2/sharing/list_shared_links",
            headers={
                "Authorization": f"Bearer {DROPBOX_TOKEN}",
                "Content-Type": "application/json",
            },
            json={"path": dropbox_path, "direct_only": True},
        )
        resp2.raise_for_status()
        links = resp2.json().get("links", [])
        if links:
            return links[0]["url"]
        raise RuntimeError(f"No shared links found for {dropbox_path}")
    elif error:
        raise RuntimeError(f"Could not create shared link: {error}")
    else:
        return data["url"]


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# Blog + RSS feed updates
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

def make_slug(show_name):
    """Create a URL-friendly slug from a show name."""
    return re.sub(r'[^a-z0-9]+', '-', show_name.lower()).strip('-')


def html_escape(text):
    """Basic HTML escaping."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def update_blog(episode, summary, dropbox_links):
    """Add a new episode entry to the allEpisodes array in index.html."""
    today = datetime.now().strftime("%B %d, %Y")
    date_slug = datetime.now().strftime("%Y-%m-%d")
    show_slug = make_slug(episode["show_name"])
    entry_id = f"{date_slug}-{show_slug}"

    km = summary.get("key_moment", {})
    km_ts = km.get("timestamp_mm_ss", "0:00")
    km_secs = km.get("timestamp_seconds", 0)
    km_text = km.get("text", "")
    km_speaker = km.get("speaker", "Unknown")

    sentiment = summary.get("sentiment", "neutral").title()
    appearance = summary.get("appearance_type", episode.get("match_type", "mentioned"))
    match_label = "Guest Appearance" if appearance == "guest" else "Mentioned"
    source_label = "Video" if episode.get("video_id") else "Podcast"

    # Build tags array
    tags = [episode["search_target"], match_label, source_label]

    # Build links array
    episode_link_label = "Watch Episode" if episode.get("video_id") else "Listen to Episode"
    links = [[episode_link_label, episode["url"]]]
    if dropbox_links.get("summary") and dropbox_links["summary"] != "#":
        links.append(["Summary", dropbox_links["summary"]])
    if dropbox_links.get("transcript") and dropbox_links["transcript"] != "#":
        links.append(["Full Transcript", dropbox_links["transcript"]])

    # Build quotes array from key_quotes (limit 8)
    key_quotes = summary.get("key_quotes", [])[:8]
    quotes = []
    for kq in key_quotes:
        q_text = kq.get("text", "")
        q_speaker = kq.get("speaker", "Unknown")
        q_ts = kq.get("timestamp_mm_ss", "0:00")
        q_secs = kq.get("timestamp_seconds", 0)
        q_link = f"{episode['url']}&t={q_secs}" if q_secs else ""
        quotes.append({
            "text": q_text,
            "attribution": f"-- {q_speaker} at {q_ts}",
            "link": q_link,
        })

    # If no key_quotes, fall back to key_moment
    if not quotes and km_text:
        quote_link = f"{episode['url']}&t={km_secs}" if km_secs else ""
        quotes.append({
            "text": km_text,
            "attribution": f"-- {km_speaker} at {km_ts}",
            "link": quote_link,
        })

    # Build the JS object as a JSON string
    ep_obj = {
        "title": episode["title"],
        "show": episode["show_name"],
        "description": summary.get("overview", ""),
        "date": today,
        "tags": tags,
        "sentiment": sentiment,
        "quotes": quotes,
        "links": links,
    }

    # Serialize to JS-safe JSON (single entry to prepend to array)
    import json as _json
    ep_json = _json.dumps(ep_obj, ensure_ascii=False)

    # Read current index.html and insert into the allEpisodes array
    content = INDEX_FILE.read_text()

    insertion_point = "const allEpisodes = ["
    if insertion_point in content:
        # Insert the new episode object right after the opening bracket
        content = content.replace(
            insertion_point,
            insertion_point + "\n" + ep_json + ","
        )
    else:
        print("  WARNING: Could not find allEpisodes array in index.html")
        return

    INDEX_FILE.write_text(content)
    print(f"  Blog updated with entry: {entry_id}")
    return entry_id

def update_rss(episode, summary, dropbox_links, entry_id):
    """Add a new item to feed.xml."""
    now = datetime.now(timezone.utc)
    pub_date = now.strftime("%a, %d %b %Y %H:%M:%S +0000")

    km = summary.get("key_moment", {})
    km_ts = km.get("timestamp_mm_ss", "0:00")
    km_secs = km.get("timestamp_seconds", 0)

    headline = f"{episode['show_name'].split(' with ')[0] if ' with ' in episode['show_name'] else episode['show_name']}: {summary.get('overview', episode['title'])[:80]}"

    # Build document links for RSS – only include if Dropbox succeeded
    doc_links_html = ""
    if dropbox_links.get("summary") and dropbox_links["summary"] != "#":
        doc_links_html += f'\n<a href="{dropbox_links["summary"]}">📄 Summary</a>'
    if dropbox_links.get("transcript") and dropbox_links["transcript"] != "#":
        doc_links_html += f' |\n<a href="{dropbox_links["transcript"]}">📄 Full Transcript</a>'

    description_html = f"""<p><strong>{html_escape(episode['show_name'])}</strong><br/>
"{html_escape(episode['title'])}" – Published {episode['date_published']}</p>
<p><strong>Match:</strong> {html_escape(episode['search_target'])} ({episode['match_type']})</p>
<p><strong>Sentiment: {summary.get('sentiment', 'neutral').title()}</strong></p>
<p>{html_escape(summary.get('overview', ''))}</p>
<p><strong>Key moment at <a href="{episode['url']}&t={km_secs}">{km_ts}</a>:</strong> "{html_escape(km.get('text', ''))}" – {html_escape(km.get('speaker', 'Unknown'))}</p>
<p><a href="{episode['url']}">▶ {"Watch Episode" if episode.get("video_id") else "Listen to Episode"}</a>{doc_links_html}</p>"""

    item_xml = f"""    <item>
      <title>{html_escape(headline)}</title>
      <link>{BLOG_URL}#{entry_id}</link>
      <guid>{BLOG_URL}#{entry_id}</guid>
      <pubDate>{pub_date}</pubDate>
      <description><![CDATA[{description_html}]]></description>
    </item>"""

    content = FEED_FILE.read_text()

    # Insert after the <atom:link> line
    atom_link_pattern = r'(<atom:link[^/]*/>\s*)'
    match = re.search(atom_link_pattern, content)
    if match:
        insert_pos = match.end()
        content = content[:insert_pos] + "\n" + item_xml + "\n" + content[insert_pos:]
    else:
        # Fallback: insert before </channel>
        content = content.replace("</channel>", item_xml + "\n  </channel>")

    FEED_FILE.write_text(content)
    print(f"  RSS feed updated")


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# Knowledge base (local transcript store)
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

def save_to_knowledge_base(episode, transcript_data, summary):
    """Save transcript + summary as a searchable .md file in the kb/ folder."""
    KB_DIR.mkdir(parents=True, exist_ok=True)

    date_str = datetime.now().strftime("%Y-%m-%d")
    safe_title = re.sub(r'[^\w\s-]', '', episode['title']).strip()
    safe_show = re.sub(r'[^\w\s-]', '', episode['show_name']).strip()
    filename = f"{date_str}--{safe_show}--{safe_title}.md"
    # Keep filename length reasonable
    if len(filename) > 200:
        filename = filename[:196] + ".md"

    lines = []
    lines.append(f"# {episode['title']}")
    lines.append("")
    lines.append(f"**Show:** {episode['show_name']}")
    lines.append(f"**Date Published:** {episode['date_published']}")
    lines.append(f"**Date Found:** {date_str}")
    lines.append(f"**URL:** {episode['url']}")
    lines.append(f"**Search Target:** {episode['search_target']} ({episode['match_type']})")
    duration_min = transcript_data.get('audio_duration', 0) // 60
    lines.append(f"**Duration:** {duration_min} minutes")
    lines.append("")

    # Summary section
    lines.append("## Summary")
    lines.append("")
    lines.append(summary.get("overview", ""))
    lines.append("")
    lines.append(f"**Sentiment:** {summary.get('sentiment', 'unknown')}")
    lines.append(f"**Sentiment Detail:** {summary.get('sentiment_explanation', '')}")
    lines.append("")

    topics = summary.get("topics", [])
    if topics:
        lines.append(f"**Topics:** {', '.join(topics)}")
        lines.append("")

    # Key moment
    km = summary.get("key_moment", {})
    if km:
        lines.append("## Key Moment")
        lines.append("")
        lines.append(f"> \"{km.get('text', '')}\"")
        lines.append(f"> â {km.get('speaker', 'Unknown')} at {km.get('timestamp_mm_ss', '??:??')}")
        if km.get("context"):
            lines.append(f"> *{km['context']}*")
        lines.append("")

    # Key quotes
    key_quotes = summary.get("key_quotes", [])
    if key_quotes:
        lines.append("## Key Quotes")
        lines.append("")
        for q in key_quotes:
            ts = q.get("timestamp_mm_ss", "??:??")
            speaker = q.get("speaker", "Unknown")
            lines.append(f"- [{ts}] {speaker}: \"{q['text']}\"")
        lines.append("")

    # Full transcript
    lines.append("## Full Transcript")
    lines.append("")
    for u in transcript_data.get("utterances", []):
        start_sec = u["start"] // 1000
        mm, ss = divmod(start_sec, 60)
        lines.append(f"[{mm:02d}:{ss:02d}] **Speaker {u['speaker']}:** {u['text']}")
        lines.append("")

    kb_path = KB_DIR / filename
    kb_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"  Knowledge base: saved {filename} ({len(lines)} lines)")
    return str(kb_path)


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# Main pipeline
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

def process_episode(episode, episodes_data):
    """Full pipeline for a single episode."""
    print(f"\n{'='*60}")
    print(f"Processing: {episode['title']}")
    print(f"  Show: {episode['show_name']}")
    print(f"  Target: {episode['search_target']} ({episode['match_type']})")
    print(f"{'='*60}")

    # 1. Transcribe
    transcript_data = transcribe_episode(episode)

    if not transcript_data:
        print("  WARNING: All transcription methods failed for this episode")
        print("  Recording episode as found (without transcript)")
        episode["status"] = "transcript_unavailable"
        episode["processed_date"] = datetime.now().isoformat()
        episodes_data["episodes"].append(episode)
        save_episodes(episodes_data)
        print(f"  Episode logged: {episode['title']} (transcript unavailable)")
        return True

    # 2. Generate summary
    print("Generating summary...")
    summary = generate_summary_with_claude(episode, transcript_data)

    # 3. Save to knowledge base
    save_to_knowledge_base(episode, transcript_data, summary)

    # 4. Create .docx files
    safe_title = re.sub(r'[^\w\s-]', '', episode['title']).strip()
    safe_show = re.sub(r'[^\w\s-]', '', episode['show_name']).strip()
    base_name = f"{safe_show} - {safe_title}"

    tmpdir = tempfile.mkdtemp()
    summary_path = os.path.join(tmpdir, f"{base_name} - Summary.docx")
    transcript_path = os.path.join(tmpdir, f"{base_name} - Transcript.docx")

    create_summary_docx(episode, summary, summary_path)
    create_transcript_docx(episode, transcript_data, transcript_path)

    # 5. Upload to Dropbox (non-fatal – blog still updates if this fails)
    dbx_summary = f"/{base_name} - Summary.docx"
    dbx_transcript = f"/{base_name} - Transcript.docx"

    dropbox_links = {"summary": "#", "transcript": "#"}
    dropbox_paths = {"summary": dbx_summary, "transcript": dbx_transcript}

    try:
        upload_to_dropbox(summary_path, dbx_summary)
        upload_to_dropbox(transcript_path, dbx_transcript)

        # 6. Get shared links
        summary_link = get_dropbox_link(dbx_summary)
        transcript_link = get_dropbox_link(dbx_transcript)
        print(f"  Summary link: {summary_link}")
        print(f"  Transcript link: {transcript_link}")
        dropbox_links = {"summary": summary_link, "transcript": transcript_link}
    except Exception as e:
        print(f"  WARNING: Dropbox upload/link failed: {e}")
        print("  Continuing without Dropbox links – blog will still be updated")

    # 7. (Blog/RSS now updated as digest after all episodes processed)

    # 8. Update episode log
    episodes_data["episodes"].append({
        "url": episode["url"],
        "title": episode["title"],
        "show_name": episode["show_name"],
        "date_found": datetime.now().strftime("%Y-%m-%d"),
        "date_published": episode["date_published"],
        "search_target": episode["search_target"],
        "match_type": episode["match_type"],
        "dropbox_paths": dropbox_paths,
        "dropbox_links": dropbox_links,
    })
    save_episodes(episodes_data)

    # Cleanup
    try:
        os.unlink(summary_path)
        os.unlink(transcript_path)
        os.rmdir(tmpdir)
    except OSError:
        pass

    print(f"\nDone processing: {episode['title']}")
    return {"episode": episode, "summary": summary, "dropbox_links": dropbox_links}



def update_blog_links(episode, dropbox_links):
    """Update an existing blog entry to add Dropbox document links."""
    if not INDEX_FILE.exists():
        return
    content = INDEX_FILE.read_text()

    # Find the digest entry for this episode by matching the episode URL
    ep_url = episode["url"]
    link_label = "Watch Episode" if episode.get("video_id") else "Listen to Episode"
    watch_link = f'<a href="{ep_url}">&#9654; {link_label}</a>'
    if watch_link not in content:
        escaped_url = html_escape(ep_url)
        watch_link = f'<a href="{escaped_url}">&#9654; {link_label}</a>'
    if watch_link not in content:
        print(f"  WARNING: Could not find blog entry to update for: {episode['title']}")
        return

    # Build new doc links
    doc_links = ""
    if dropbox_links.get("summary") and dropbox_links["summary"] != "#":
        doc_links += f""" |
        <a href="{html_escape(dropbox_links['summary'])}">&#128196; Summary</a>"""
    if dropbox_links.get("transcript") and dropbox_links["transcript"] != "#":
        doc_links += f""" |
        <a href="{html_escape(dropbox_links['transcript'])}">&#128196; Full Transcript</a>"""

    if doc_links:
        content = content.replace(watch_link, watch_link + doc_links)
        INDEX_FILE.write_text(content)
        print(f"  Blog entry updated with Dropbox links")

def reprocess_failed_episodes(episodes_data):
    """Retry failed episodes: re-transcribe if needed, re-upload to Dropbox, update blog."""
    episodes = episodes_data["episodes"]
    to_reprocess = []

    for i, ep in enumerate(episodes):
        needs_work = False
        reason = ""
        if ep.get("status") == "transcript_unavailable":
            needs_work = True
            reason = "transcript_unavailable"
        elif ep.get("dropbox_links", {}).get("summary") == "#":
            needs_work = True
            reason = "dropbox_failed"
        if needs_work:
            to_reprocess.append((i, ep, reason))

    if not to_reprocess:
        print("\nNo failed episodes to reprocess.")
        return

    print(f"\nFound {len(to_reprocess)} episode(s) to reprocess")

    processed = 0
    for idx, ep, reason in to_reprocess:
        try:
            print(f"\n{'='*60}")
            print(f"Reprocessing: {ep['title']}")
            print(f"  Reason: {reason}")
            print(f"{'='*60}")

            # Build episode dict
            episode = {
                "url": ep["url"],
                "video_id": ep.get("video_id", ""),
                "title": ep["title"],
                "show_name": ep["show_name"],
                "date_published": ep.get("date_published", ""),
                "search_target": ep["search_target"],
                "match_type": ep["match_type"],
            }
            if ep.get("audio_url"):
                episode["audio_url"] = ep["audio_url"]
            # Extract video_id from URL if not present
            if not episode["video_id"] and "youtube.com/watch" in episode["url"]:
                import urllib.parse
                parsed = urllib.parse.urlparse(episode["url"])
                qs = urllib.parse.parse_qs(parsed.query)
                episode["video_id"] = qs.get("v", [""])[0]
            # For non-YouTube episodes without audio_url, use the episode URL directly
            if not episode.get("video_id") and not episode.get("audio_url"):
                episode["audio_url"] = episode["url"]

            # Try to load existing transcript from knowledge base
            transcript_data = None
            safe_title_kb = re.sub(r'[^\w\s-]', '', episode['title']).strip()
            safe_show_kb = re.sub(r'[^\w\s-]', '', episode['show_name']).strip()
            kb_pattern = f"*--{safe_show_kb}--{safe_title_kb}*"
            kb_matches = list(KB_DIR.glob(kb_pattern))
            if kb_matches:
                kb_text = kb_matches[0].read_text(encoding="utf-8")
                # Extract transcript section
                tx_marker = "## Full Transcript"
                if tx_marker in kb_text:
                    tx_section = kb_text[kb_text.index(tx_marker) + len(tx_marker):].strip()
                    utterances = []
                    full_text_parts = []
                    for line in tx_section.split("\n"):
                        m = re.match(r'\[(\d+):(\d+)\] \*\*Speaker (\w+):\*\* (.+)', line)
                        if m:
                            mm, ss, speaker, text = m.groups()
                            start_ms = (int(mm) * 60 + int(ss)) * 1000
                            utterances.append({"start": start_ms, "end": start_ms + 1000, "text": text, "speaker": speaker})
                            full_text_parts.append(text)
                    if utterances:
                        transcript_data = {
                            "utterances": utterances,
                            "text": " ".join(full_text_parts),
                            "audio_duration": utterances[-1]["start"] // 1000,
                            "source": "kb_file",
                        }
                        print(f"  Loaded existing transcript from KB ({len(utterances)} segments)")

            # If no KB transcript, try fresh transcription
            if not transcript_data:
                transcript_data = transcribe_episode(episode)
            if not transcript_data:
                print("  Still unable to transcribe. Skipping.")
                continue

            # Generate summary
            summary = generate_summary_with_claude(episode, transcript_data)

            # Save to knowledge base
            save_to_knowledge_base(episode, transcript_data, summary)

            # Create docx files
            safe_title = re.sub(r'[^\w\s-]', '', episode['title']).strip()
            safe_show = re.sub(r'[^\w\s-]', '', episode['show_name']).strip()
            base_name = f"{safe_show} - {safe_title}"
            tmpdir = tempfile.mkdtemp()
            summary_path = os.path.join(tmpdir, f"{base_name} - Summary.docx")
            transcript_path = os.path.join(tmpdir, f"{base_name} - Transcript.docx")
            create_summary_docx(episode, summary, summary_path)
            create_transcript_docx(episode, transcript_data, transcript_path)

            # Upload to Dropbox
            dbx_summary = f"/{base_name} - Summary.docx"
            dbx_transcript = f"/{base_name} - Transcript.docx"
            dropbox_links = {"summary": "#", "transcript": "#"}

            try:
                upload_to_dropbox(summary_path, dbx_summary)
                upload_to_dropbox(transcript_path, dbx_transcript)
                summary_link = get_dropbox_link(dbx_summary)
                transcript_link = get_dropbox_link(dbx_transcript)
                dropbox_links = {"summary": summary_link, "transcript": transcript_link}
                print(f"  Summary link: {summary_link}")
                print(f"  Transcript link: {transcript_link}")
            except Exception as e:
                print(f"  WARNING: Dropbox upload failed: {e}")

            # Update or add blog entry
            if reason == "transcript_unavailable":
                # This episode never got a blog entry, so add one
                entry_id = update_blog(episode, summary, dropbox_links)
                if entry_id:
                    update_rss(episode, summary, dropbox_links, entry_id)
            elif transcript_data.get("source") == "kb_file":
                # Had a transcript but bad summary — replace the blog entry
                entry_id = update_blog(episode, summary, dropbox_links)
                if entry_id:
                    update_rss(episode, summary, dropbox_links, entry_id)
            else:
                # Blog entry exists, just update the Dropbox links
                update_blog_links(episode, dropbox_links)

            # Update episode record in place
            episodes_data["episodes"][idx] = {
                "url": episode["url"],
                "title": episode["title"],
                "show_name": episode["show_name"],
                "date_found": ep.get("date_found", datetime.now().strftime("%Y-%m-%d")),
                "date_published": episode["date_published"],
                "search_target": episode["search_target"],
                "match_type": episode["match_type"],
                "dropbox_paths": {"summary": dbx_summary, "transcript": dbx_transcript},
                "dropbox_links": dropbox_links,
            }
            save_episodes(episodes_data)

            # Cleanup
            try:
                os.unlink(summary_path)
                os.unlink(transcript_path)
                os.rmdir(tmpdir)
            except OSError:
                pass

            processed += 1

        except Exception as e:
            print(f"\nERROR reprocessing {ep['title']}: {e}")
            import traceback
            traceback.print_exc()
            continue

    print(f"\n{'='*60}")
    print(f"Reprocess complete. Fixed {processed}/{len(to_reprocess)} episodes.")
    print(f"{'='*60}")


def update_digest(digest_results):
    """Add a single digest entry to index.html containing all episodes from this scan."""
    today = datetime.now().strftime("%B %d, %Y")
    date_slug = datetime.now().strftime("%Y-%m-%d")
    time_slug = datetime.now().strftime("%H%M")
    entry_id = f"{date_slug}-digest-{time_slug}"

    episode_count = len(digest_results)

    # Build HTML for each episode within the digest
    episodes_html = ""
    for result in digest_results:
        episode = result["episode"]
        summary = result["summary"]
        dropbox_links = result["dropbox_links"]

        sentiment_class = f"sentiment-{summary.get('sentiment', 'neutral')}"
        appearance = summary.get("appearance_type", episode.get("match_type", "mentioned"))
        match_class = "tag-guest" if appearance == "guest" else "tag-mention"
        match_label = "Guest Appearance" if appearance == "guest" else "Mentioned"
        source_label = "Video" if episode.get("video_id") else "Podcast"

        doc_links = ""
        if dropbox_links.get("summary") and dropbox_links["summary"] != "#":
            doc_links += f' |\n            <a href="{html_escape(dropbox_links["summary"])}">&#128196; Summary</a>'
        if dropbox_links.get("transcript") and dropbox_links["transcript"] != "#":
            doc_links += f' |\n            <a href="{html_escape(dropbox_links["transcript"])}">&#128196; Full Transcript</a>'

        # Build multiple quotes HTML (limit 8)
        key_quotes = summary.get("key_quotes", [])[:8]
        km = summary.get("key_moment", {})
        if not key_quotes and km.get("text"):
            key_quotes = [km]

        quotes_html = ""
        for kq in key_quotes:
            q_text = html_escape(kq.get("text", ""))
            q_speaker = html_escape(kq.get("speaker", "Unknown"))
            q_ts = kq.get("timestamp_mm_ss", "0:00")
            q_secs = kq.get("timestamp_seconds", 0)
            quotes_html += f"""
          <p>&ldquo;{q_text}&rdquo;</p>
          <cite>&mdash; {q_speaker} at <a href="{episode['url']}&amp;t={q_secs}">{q_ts}</a></cite>"""

        episodes_html += f"""
      <div class="digest-episode">
        <h3>{html_escape(episode['show_name'])}</h3>
        <h4>{html_escape(episode['title'])}</h4>
        <div class="tags">
          <span class="tag">{html_escape(episode['search_target'])}</span>
          <span class="tag {match_class}">{match_label}</span>
          <span class="tag tag-source">{source_label}</span>
        </div>
        <div class="sentiment {sentiment_class}">Sentiment: {summary.get('sentiment', 'neutral').title()}</div>
        <p>{html_escape(summary.get('overview', ''))}</p>
        <blockquote class="quote">{quotes_html}
        </blockquote>
        <div class="links">
          <a href="{episode['url']}">&#9654; {"Watch Episode" if episode.get("video_id") else "Listen to Episode"}</a>{doc_links}
        </div>
      </div>"""

    digest_html = f"""    <div class="digest" id="{entry_id}">
      <div class="digest-date">{today}</div>
      <h2>PYT Radar Digest &mdash; {today}</h2>
      <p class="digest-summary">{episode_count} new episode{'s' if episode_count != 1 else ''} found matching tracked targets.</p>{episodes_html}
    </div>"""

    content = INDEX_FILE.read_text()
    content = re.sub(r'<p class="empty">.*?</p>\s*', '', content)

    insertion_point = '<main id="digests">'
    if insertion_point in content:
        content = content.replace(insertion_point, insertion_point + "\n" + digest_html)
    else:
        print("  WARNING: Could not find insertion point in index.html")
        return None

    INDEX_FILE.write_text(content)
    print(f"  Blog updated with digest: {entry_id} ({episode_count} episodes)")
    return entry_id


def update_digest_rss(digest_results, entry_id):
    """Add a single RSS item for the digest containing all episodes."""
    now = datetime.now(timezone.utc)
    pub_date = now.strftime("%a, %d %b %Y %H:%M:%S +0000")
    today = datetime.now().strftime("%B %d, %Y")

    episode_count = len(digest_results)

    if episode_count == 1:
        ep = digest_results[0]["episode"]
        sm = digest_results[0]["summary"]
        headline = f"{ep['show_name']}: {sm.get('overview', ep['title'])[:80]}"
    else:
        targets = set()
        for r in digest_results:
            targets.add(r["episode"]["search_target"])
        headline = f"PYT Radar: {episode_count} new episodes ({', '.join(targets)})"

    episodes_desc = ""
    for result in digest_results:
        episode = result["episode"]
        summary = result["summary"]
        dropbox_links = result["dropbox_links"]

        km = summary.get("key_moment", {})
        km_ts = km.get("timestamp_mm_ss", "0:00")
        km_secs = km.get("timestamp_seconds", 0)

        doc_links_html = ""
        if dropbox_links.get("summary") and dropbox_links["summary"] != "#":
            doc_links_html += f'\n<a href="{dropbox_links["summary"]}">&#128196; Summary</a>'
        if dropbox_links.get("transcript") and dropbox_links["transcript"] != "#":
            doc_links_html += f' |\n<a href="{dropbox_links["transcript"]}">&#128196; Full Transcript</a>'

        episodes_desc += f"""<hr/>
<p><strong>{html_escape(episode['show_name'])}</strong><br/>
\"{html_escape(episode['title'])}\" \u2013 Published {episode['date_published']}</p>
<p><strong>Match:</strong> {html_escape(episode['search_target'])} ({episode['match_type']})</p>
<p><strong>Sentiment: {summary.get('sentiment', 'neutral').title()}</strong></p>
<p>{html_escape(summary.get('overview', ''))}</p>
<p><strong>Key moment at <a href="{episode['url']}&t={km_secs}">{km_ts}</a>:</strong> \"{html_escape(km.get('text', ''))}\" \u2013 {html_escape(km.get('speaker', 'Unknown'))}</p>
<p><a href="{episode['url']}">&#9654; {"Watch Episode" if episode.get("video_id") else "Listen to Episode"}</a>{doc_links_html}</p>"""

    description_html = f"""<h2>PYT Radar Digest \u2014 {today}</h2>
<p>{episode_count} new episode{'s' if episode_count != 1 else ''} found.</p>
{episodes_desc}"""

    item_xml = f"""    <item>
      <title>{html_escape(headline)}</title>
      <link>{BLOG_URL}#{entry_id}</link>
      <guid>{BLOG_URL}#{entry_id}</guid>
      <pubDate>{pub_date}</pubDate>
      <description><![CDATA[{description_html}]]></description>
    </item>"""

    content = FEED_FILE.read_text()
    atom_link_pattern = r'(<atom:link[^/]*/>[\\s]*)'
    match_obj = re.search(atom_link_pattern, content)
    if match_obj:
        insert_pos = match_obj.end()
        content = content[:insert_pos] + "\n" + item_xml + "\n" + content[insert_pos:]
    else:
        content = content.replace("</channel>", item_xml + "\n  </channel>")

    FEED_FILE.write_text(content)
    print(f"  RSS feed updated with digest")


def send_cookie_expiry_alert():
    """Send alert email when YouTube cookies are detected as expired."""
    mailgun_api_key = os.environ.get("MAILGUN_API_KEY")
    mailgun_domain = os.environ.get("MAILGUN_DOMAIN")
    if not mailgun_api_key or not mailgun_domain:
        print("  Mailgun not configured, skipping cookie alert")
        return

    config_path = REPO_ROOT / "references" / "config.json"
    if not config_path.exists():
        return
    config = json.loads(config_path.read_text())
    recipients = config.get("recipients", [])
    if not recipients:
        return

    today = datetime.now().strftime("%B %d, %Y %I:%M %p")
    email_html = f"""<!DOCTYPE html>
<html>
<body style="background: #0a0a0a; color: #e0e0e0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; margin: 0; padding: 0;">
  <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
    <div style="background: #2a1111; border: 1px solid #d44; padding: 20px; border-radius: 10px;">
      <h2 style="color: #d44; margin: 0 0 8px 0;">&#9888; PYT Radar: YouTube Cookies Expired</h2>
      <p style="color: #ccc; margin: 0 0 12px 0;">Detected at {today}</p>
      <p style="color: #e0e0e0;">YouTube is blocking video downloads because the authentication cookies have expired or been invalidated. Episodes found via YouTube captions will still work, but any video without auto-captions will be skipped until cookies are refreshed.</p>
      <p style="color: #e0e0e0; margin-top: 12px;"><strong>To fix:</strong></p>
      <ol style="color: #ccc;">
        <li>Open YouTube in Chrome and make sure you're signed in</li>
        <li>Export cookies using a browser extension (e.g. "Get cookies.txt LOCALLY")</li>
        <li>Update the YOUTUBE_COOKIES secret in the <a href="https://github.com/ericspiegelman/pyt-radar/settings/secrets/actions" style="color: #6cb4ee;">pyt-radar repo settings</a></li>
      </ol>
    </div>
  </div>
</body>
</html>"""

    try:
        resp = requests.post(
            f"https://api.mailgun.net/v3/{mailgun_domain}/messages",
            auth=("api", mailgun_api_key),
            data={{
                "from": f"PYT Radar <pyt-radar@{mailgun_domain}>",
                "to": recipients,
                "subject": "PYT Radar Alert: YouTube cookies expired",
                "html": email_html,
            }},
        )
        if resp.status_code == 200:
            print("  Cookie expiry alert sent")
        else:
            print(f"  WARNING: Cookie alert email failed: {resp.status_code}")
    except Exception as e:
        print(f"  WARNING: Failed to send cookie alert: {e}")


def send_digest_email(digest_results):
    """Send digest email via Mailgun to all recipients in config."""
    mailgun_api_key = os.environ.get("MAILGUN_API_KEY")
    mailgun_domain = os.environ.get("MAILGUN_DOMAIN")

    if not mailgun_api_key or not mailgun_domain:
        print("  Mailgun not configured, skipping email delivery")
        return

    # Load recipients from config
    config_path = REPO_ROOT / "references" / "config.json"
    if not config_path.exists():
        print("  No config.json found, skipping email")
        return
    config = json.loads(config_path.read_text())
    recipients = config.get("recipients", [])
    if not recipients:
        print("  No recipients configured, skipping email")
        return

    today = datetime.now().strftime("%B %d, %Y")
    episode_count = len(digest_results)

    # Build email HTML
    episodes_html = ""
    for result in digest_results:
        episode = result["episode"]
        summary = result["summary"]
        dropbox_links = result["dropbox_links"]

        doc_links_html = ""
        if dropbox_links.get("summary") and dropbox_links["summary"] != "#":
            doc_links_html += f'<a href="{dropbox_links["summary"]}" style="color: #3b82f6; text-decoration: none; font-weight: 500;">Summary</a>'
        if dropbox_links.get("transcript") and dropbox_links["transcript"] != "#":
            doc_links_html += f' &nbsp;&middot;&nbsp; <a href="{dropbox_links["transcript"]}" style="color: #3b82f6; text-decoration: none; font-weight: 500;">Full Transcript</a>'

        sentiment = summary.get("sentiment", "neutral").title()
        sentiment_colors = {"Positive": "#50c080", "Negative": "#e05050", "Neutral": "#888", "Mixed": "#e0a050"}
        s_color = sentiment_colors.get(sentiment, "#888")

        appearance = summary.get("appearance_type", episode.get("match_type", "mentioned"))
        appearance_label = "Guest Appearance" if appearance == "guest" else "Mentioned"
        source_label = "Video" if episode.get("video_id") else "Podcast"

        # Build quotes HTML (multiple quotes, limit 8)
        key_quotes = summary.get("key_quotes", [])[:8]
        km = summary.get("key_moment", {})
        if not key_quotes and km.get("text"):
            key_quotes = [km]

        quotes_html = ""
        for kq in key_quotes:
            q_text = kq.get("text", "")
            if len(q_text) > 200:
                q_text = q_text[:200].rsplit(" ", 1)[0] + "..."
            q_speaker = html_escape(kq.get("speaker", "Unknown"))
            q_ts = kq.get("timestamp_mm_ss", "0:00")
            q_secs = kq.get("timestamp_seconds", 0)
            quotes_html += f"""
            <div style="padding: 8px 0; border-bottom: 1px solid #f0f0f0;">
              <p style="color: #555; font-style: italic; margin: 0 0 4px 0; font-size: 13px; line-height: 1.5;">&ldquo;{html_escape(q_text)}&rdquo;</p>
              <cite style="color: #999; font-size: 12px; font-style: normal;">&mdash; {q_speaker} at <a href="{episode['url']}&t={q_secs}" style="color: #3b82f6; text-decoration: none;">{q_ts}</a></cite>
            </div>"""

        episodes_html += f"""
        <div style="background: #ffffff; border: 1px solid #e8e8e8; border-radius: 8px; padding: 20px; margin-top: 16px;">
          <p style="color: #999; font-size: 11px; text-transform: uppercase; letter-spacing: 0.5px; margin: 0 0 4px 0; font-weight: 600;">{html_escape(episode['show_name'])}</p>
          <h3 style="color: #1a1a1a; margin: 0 0 10px 0; font-size: 16px; font-weight: 600;">{html_escape(episode['title'])}</h3>
          <p style="margin: 0 0 10px 0;">
            <span style="background: #eef4fb; color: #3b82f6; padding: 3px 10px; border-radius: 12px; font-size: 12px; font-weight: 500;">{html_escape(episode['search_target'])}</span>
            <span style="background: #f0e8f4; color: #7b1fa2; padding: 3px 10px; border-radius: 12px; font-size: 12px; font-weight: 500; margin-left: 4px;">{appearance_label}</span>
            <span style="background: #e8f5e9; color: #2e7d32; padding: 3px 10px; border-radius: 12px; font-size: 12px; font-weight: 500; margin-left: 4px;">{source_label}</span>
            <span style="background: {'#fef2f2' if sentiment == 'Negative' else '#f8f8f8'}; color: {s_color}; padding: 3px 10px; border-radius: 12px; font-size: 12px; font-weight: 500; margin-left: 4px;">{sentiment}</span>
          </p>
          <p style="color: #555; margin: 10px 0; font-size: 14px; line-height: 1.5;">{html_escape(summary.get('overview', ''))}</p>
          <blockquote style="border-left: 3px solid #e74c3c; padding: 10px 14px; margin: 12px 0; background: #fafafa; border-radius: 0 6px 6px 0;">
            {quotes_html}
          </blockquote>
          <p style="margin: 12px 0 0 0; font-size: 13px;">
            <a href="{episode['url']}" style="color: #3b82f6; text-decoration: none; font-weight: 500;">&#9654; {"Watch Episode" if episode.get("video_id") else "Listen to Episode"}</a>
            {(' &nbsp;&middot;&nbsp; ' + doc_links_html) if doc_links_html else ''}
          </p>
        </div>"""

    email_html = f"""<!DOCTYPE html>
<html>
<body style="background: #f5f5f5; color: #2c3e50; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; margin: 0; padding: 0;">
  <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
    <div style="background: #ffffff; border-bottom: 1px solid #e0e0e0; padding: 24px; text-align: center; border-radius: 10px 10px 0 0;">
      <h1 style="color: #1a1a1a; margin: 0 0 6px 0; font-size: 24px; font-weight: 700; letter-spacing: -0.5px;">PYT Radar</h1>
      <p style="color: #888; margin: 0; font-size: 14px;">Podcast &amp; YouTube Tracking Digest - Los Angeles 2026 Mayoral Candidates</p>
    </div>
    <div style="background: #fafafa; border: 1px solid #e0e0e0; border-top: none; padding: 20px; border-radius: 0 0 10px 10px;">
      <h2 style="color: #1a1a1a; margin: 0 0 4px 0; font-size: 18px; font-weight: 600;">{today}</h2>
      <p style="color: #888; margin: 0 0 8px 0; font-size: 13px;">{episode_count} new episode{'s' if episode_count != 1 else ''} found matching tracked targets.</p>
      {episodes_html}
    </div>
    <p style="text-align: center; color: #999; font-size: 12px; margin-top: 16px;">
      <a href="{BLOG_URL}" style="color: #3b82f6; text-decoration: none;">View all digests on the web</a>
    </p>
  </div>
</body>
</html>"""

    # Send via Mailgun
    try:
        resp = requests.post(
            f"https://api.mailgun.net/v3/{mailgun_domain}/messages",
            auth=("api", mailgun_api_key),
            data={
                "from": f"PYT Radar <pyt-radar@{mailgun_domain}>",
                "to": recipients,
                "subject": f"PYT Radar: {episode_count} new episode{'s' if episode_count != 1 else ''} - {today}",
                "html": email_html,
            },
        )
        if resp.status_code == 200:
            print(f"  Digest email sent to {len(recipients)} recipient(s)")
        else:
            print(f"  WARNING: Mailgun returned {resp.status_code}: {resp.text}")
    except Exception as e:
        print(f"  WARNING: Failed to send digest email: {e}")


def main():
    parser = argparse.ArgumentParser(description="PYT Radar scanner")
    parser.add_argument("--mode", choices=["youtube", "podcast", "all"], default="all",
                        help="Search mode: youtube, podcast, or all (default: all)")
    parser.add_argument("--reprocess", action="store_true",
                        help="Reprocess failed episodes (retry transcription and Dropbox uploads)")
    args = parser.parse_args()

    print("=" * 60)
    print(f"PYT Radar â Scan starting (mode: {args.mode})")
    print(f"Time: {datetime.now().isoformat()}")
    print("=" * 60)

    # Load episode log
    episodes_data = load_episodes()
    print(f"Loaded {len(episodes_data['episodes'])} previously found episodes")

    # Refresh Dropbox token (gets a fresh short-lived token from the long-lived refresh token)
    refresh_dropbox_token()

    # Handle reprocess mode
    if args.reprocess:
        reprocess_failed_episodes(episodes_data)
        return

    # Search for new episodes
    new_episodes = find_new_episodes(episodes_data, mode=args.mode)

    if not new_episodes:
        print("\nNo new episodes found. Scan complete.")
        return

    print(f"\nFound {len(new_episodes)} new episode(s) to process")

    # Process each new episode
    processed = 0
    digest_results = []
    for episode in new_episodes:
        try:
            result = process_episode(episode, episodes_data)
            if isinstance(result, dict):
                processed += 1
                digest_results.append(result)
            elif result:
                processed += 1
        except Exception as e:
            print(f"\nERROR processing {episode['title']}: {e}")
            continue

    # Create digest blog post and RSS entry for all episodes found in this scan
    if digest_results:
        entry_id = update_digest(digest_results)
        if entry_id:
            update_digest_rss(digest_results, entry_id)
        send_digest_email(digest_results)

    # Check for YouTube cookie issues and alert
    if _youtube_cookies_expired:
        print("\n  YouTube cookies expired - sending alert...")
        send_cookie_expiry_alert()

    print(f"\n{'='*60}")
    print(f"Scan complete. Processed {processed}/{len(new_episodes)} episodes.")
    print(f"{'='*60}")

    # Set output for GitHub Actions
    if processed > 0:
        github_output = os.environ.get("GITHUB_OUTPUT", "")
        if github_output:
            with open(github_output, "a") as f:
                f.write(f"episodes_found={processed}\n")


if __name__ == "__main__":
    main()
